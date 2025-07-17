import sys
tt, TT, ff, FF = True, True, False, False
from docx import Document
import warnings
warnings.filterwarnings('ignore', category=UserWarning)
import unicodedata

#-----------------------------

path_def = 'C:\\Users\\esteb\\__Files'   #home

Months = ['ene','jan','feb','mar','apr','abr','may','jun','jul','aug','ago','sep','oct','nov','dec','dic']
DatesN = ['-1-','-2-','-3-','-4-','-5-','-6-','-7-','-8-','-9-',
          '-01-','-02-','-03-','-04-','-05-','-06-','-07-','-08-','-09-',
          '-10-','-11-','-12-','-13-','-14-','-15-','-16-','-17-','-18-','-19-',
          '-20-','-21-','-22-','-23-','-24-','-25-','-26-','-27-','-28-','-29-','-30-','-31-',
          '- 1 -','- 2 -','- 3 -','- 4 -','- 5 -','- 6 -','- 7 -','- 8 -','- 9 -',
          '- 01 -','- 02 -','- 03 -','- 04 -','- 05 -','- 06 -','- 07 -','- 08 -','- 09 -',
          '- 10 -','- 11 -','- 12 -','- 13 -','- 14 -','- 15 -','- 16 -','- 17 -','- 18 -','- 19 -',
          '- 20 -','- 21 -','- 22 -','- 23 -','- 24 -','- 25 -','- 26 -','- 27 -','- 28 -','- 29 -','- 30 -','- 31 -']          
          
Styles = [' jan','jan ','-jan','-jan-','jan-','/jan','/jan/','jan/',' ene','ene ','-ene','-ene-','ene-','/ene','/ene/','ene/',
          ' feb','feb ','-feb','-feb-','feb-','/feb','/feb/','feb/',
          ' mar','mar ','-mar','-mar-','mar-','/mar','/mar/','mar/',
          ' apr','apr ','-apr','-apr-','apr-','/apr','/apr/','apr/',' abr','abr ','-abr','-abr-','abr-','/abr','/abr/','abr/',
          ' may','may ','-may','-may-','may-','/may','/may/','may/','mayo',
          ' jun','jun ','-jun','-jun-','jun-','/jun','/jun/','jun/','june',
          ' jul','jul ','-jul','-jul-','jul-','/jul','/jul/','jul/','july',
          ' aug','aug ','-aug','-aug-','aug-','/aug','/aug/','aug/',' ago','ago ','-ago','-ago-','ago-','/ago','/ago/','ago/',
          ' sep','sep ','-sep','-sep-','sep-','/sep','/sep/','sep/',           
          ' oct','oct ','-oct','-oct-','oct-','/oct','/oct/','oct/',        
          ' nov','nov ','-nov','-nov-','nov-','/nov','/nov/','nov/',
          ' dec','dec ','-dec','-dec-','dec-','/dec','/dec/','dec/',' dic','dic ','-dic','-dic-','dic-','/dic','/dic/','dic/']

Styles2 = ['january','enero','february','febrero','march','marzo',
          'april','abril','junio','julio','august','agosto','september','septiembre',           
          'october','octubre','november','noviembre','december','diciembre']

Excp_months = ['may ']

Multiphrase = Months + DatesN
Excpts = Excp_months + DatesN

Filenames_only = ff  #must be in ff for searching in context; elems in Multiphrase: are they in file names?
case_insen = tt
Files_docx, Files_xlsx, Files_pdf, Files_pptx = tt, ff, ff, ff

phrase_left = 10 #number of chars before 'phrase'
phrase_right = 10 #number of chars after 'phrase'

#-----------------------------
sep_chars = [chr(45),chr(8208),chr(8209),chr(8210),chr(8211),chr(8212),chr(8213),chr(8722),chr(92),chr(8260),chr(47),chr(92),
chr(822),chr(823),chr(824),chr(8725),chr(8726)] #- ‐ ‑ ‒ – — ― − \ ⁄ / \ ̶ ̷ ̸ ∕ ∖

def files_fromfolder(folder_):   #extract files from a folder
    import os
    Ubic_files, Name_files, Ubic_Name = [], [], []
    for root, dirs, files in os.walk(folder_):
        for file in files:
            ubic_str = os.path.join(root, file)
            Ubic_files.append(ubic_str)
    
    for direc in Ubic_files:
        new = []
        for i in range(-1, -len(direc)-1,-1):
            if direc[i]=='\\':
                break
            else:
                new.insert(0, direc[i])
        Name_files.append(''.join(new))
    Ubic_Name.append(Name_files)
    Ubic_Name.append(Ubic_files)    
    return Ubic_Name

def remove_accents(input_str):
    nfkd_form = unicodedata.normalize('NFKD', input_str)
    return "".join(c for c in nfkd_form if not unicodedata.combining(c))

def prev_next_phrase2(texto00, phrase_, inde_, phr_pre, phr_ext, case_insensitive=True):
    aa = inde_
    if phr_pre != '' and aa - phr_pre >0:
        prev_str = texto00[(aa - phr_pre):aa]
    elif phr_pre != '':
        prev_str = texto00[:aa]
    else:
        prev_str = ''
    bb = aa + len(phrase_)
    next_str = ('\u0332'.join(texto00[aa:bb])) + texto00[bb:bb+phr_ext]   #includes original phrase, underlined
    return prev_str + next_str

def find_indexs(text, substring):
    indexes = []
    start = 0
    while True:
        idx = text.find(substring, start)
        if idx == -1:
            break
        indexes.append(idx)
        start = idx + 1  # Move forward to allow overlapping matches
    return indexes


#***************** start *****************

Res, Res_adi, NotRead, NotSupport = [], [], [], []
MultiRes = [[] for i in range(len(Multiphrase))]
MultiRes_adi = [[] for i in range(len(Multiphrase))]
Multicont_total = [0 for i in range(len(Multiphrase))]
Res_namefs = []

print('insert dir path (n: default)')
opt1 = input()
if opt1.lower() == 'n':
    path00 = path_def                    
else:
    path00 = opt1


#all files from a folder:
namefs_ubicfs = files_fromfolder(path00)
namefs = namefs_ubicfs[0]   #files names only
lista = namefs_ubicfs[1]   #files ubications

if case_insen == tt:
    addcase = '(**case insen)'
else:
    addcase = ''

if Filenames_only:   #only case insen
    for num, file_name in enumerate(namefs):
        for i in range(len(Multiphrase)):
            if Multiphrase[i].lower() in file_name.lower():
                Res_namefs.append(file_name)

else:        
    for num, file_name in enumerate(lista):
        Multicont = [0 for i in range(len(Multiphrase))]
        #***** word case *****
        if file_name[-5:] == '.docx' and Files_docx:
            doc = Document(file_name)
            paragraphs = doc.paragraphs
            
            try:
                #paragraphs:
                for para in paragraphs:
                    if case_insen:
                        for i in range(len(Multiphrase)):                        
                            if Multiphrase[i].lower() in para.text.lower():  # Case-insensitive search
                                n_times = (para.text.lower()).count(Multiphrase[i].lower())  #how many times is the phrase in para
                                Multicont[i] += n_times

                                #new---------------------                            
                                indxs = find_indexs(para.text.lower(), Multiphrase[i].lower())
                                for ind in indxs:
                                    add_date = f"f_{num} " + prev_next_phrase2(para.text, Multiphrase[i], ind, phrase_left, phrase_right, case_insen) 
                                    c = 0
                                    #mini algo: searches a 3 digit word and then adds 1 char per side each time (ex.: 'dic'->'dica'->'ndica'...)
                                    for b in range(3, 7):
##                                        check00 = (para.text)[ind - c : ind + b ] #starting by left 
                                        check00 = (para.text)[ind - c : ind + b +1]  #original       
                                        if check00.lower() in Styles:
                                            if check00.lower() in Excpts:
                                                check00 = (para.text)[ind - c : ind + b +2]
                                                if check00[-1].isdigit():
                                                    MultiRes_adi[i].append(add_date)
                                                    break
                                                else:
                                                    break
                                            else:
                                                MultiRes_adi[i].append(add_date)
                                                break                                                   
                                                
                                        c += 1
                                        check00 = (para.text)[ind - c : ind + b +1]
                                        if check00.lower() in Styles:
                                            MultiRes_adi[i].append(add_date)
                                            break
                                        check01 = (para.text)[ind : ind + b]
                                        if check01.lower() in DatesN:
                                            for d in range(1,6):
                                                check01 = (para.text)[ind : ind + b +d]                                                
                                                if check01[-1].isdigit():
                                                    MultiRes_adi[i].append(add_date)
                                                    break                

                                    check02 = (para.text)[ind : ind + 5]   #searching complete names (as in Styles2)
                                    if check02.lower() in [x[:5] for x in Styles2]:
                                        MultiRes_adi[i].append(add_date)
##                                        break
                                
                    else:
                        for i in range(len(Multiphrase)): 
                            if Multiphrase[i] in para.text:  # Case sensitive search
                                n_times = (para.text).count(Multiphrase[i])
                                Multicont[i] += n_times

                                indxs = find_indexs(para.text, Multiphrase[i])
                                for ind in indxs:
                                    MultiRes_adi[i].append(add_date)

                #tables:
                for table_index, table in enumerate(doc.tables):
                    for row_index, row in enumerate(table.rows):
                        for cell_index, cell in enumerate(row.cells):
                            if case_insen:
                                for i in range(len(Multiphrase)):
                                    if Multiphrase[i].lower() in cell.text.lower():  # Case-insensitive search
                                        n_times = (cell.text.lower()).count(Multiphrase[i].lower())
                                        Multicont[i] += n_times
                                        
                                        indxs = find_indexs(cell.text.lower(), Multiphrase[i].lower())
                                        for ind in indxs:
                                            add_date = f"f_{num} " + prev_next_phrase2(cell.text, Multiphrase[i], ind, phrase_left, phrase_right, case_insen)
                                            c = 0
                                            #mini algo: searches a 3 digit word and then adds 1 char per side each time (ex.: 'dic'->'dica'->'ndica'...)
                                            for b in range(3, 7):
                                                check00 = (cell.text)[ind - c : ind + b +1]                 
                                                if check00.lower() in Styles:
                                                    if check00.lower() in Excpts:
                                                        check00 = (cell.text)[ind - c : ind + b +2]
                                                        if check00[-1].isdigit():
                                                            MultiRes_adi[i].append(add_date)
                                                            break
                                                        else:
                                                            break
                                                    else:
                                                        MultiRes_adi[i].append(add_date)
                                                        break                                                   

                                                c += 1
                                                check00 = (cell.text)[ind - c : ind + b +1]
                                                if check00.lower() in Styles:
                                                    MultiRes_adi[i].append(add_date)
                                                    break
                                                #--------****************---------------

                                                if check00.lower() in DatesN:
                                                    check00 = (cell.text)[ind : ind + b +1]
                                                    print(ind, check00) #debug
                                                    
                                                    if check00[-1].isdigit():
                                                        MultiRes_adi[i].append(add_date)
                                                        break
                                            check02 = (para.text)[ind : ind + 5]   #new
                                            if check02.lower() in [x[:5] for x in Styles2]:
                                                MultiRes_adi[i].append(add_date)
                                                break

                                   
                            else:
                                for i in range(len(Multiphrase)): 
                                    if Multiphrase[i] in cell.text:  # Case sensitive search
                                        n_times = (cell.text).count(Multiphrase[i])
                                        Multicont[i] += n_times

                                        indxs = find_indexs(cell.text, Multiphrase[i])
                                        for ind in indxs:
                                            MultiRes_adi[i].append(add_date)
                                                            
            except:
                NotRead.append(f"f_{num} {file_name}")
                continue

        else:
            NotSupport.append(f"f_{num} {file_name}")
            continue   

        for p in range(len(Multiphrase)):   
            Multicont_total[p] += Multicont[p]
            
        for j in range(len(Multicont)):
            if Multicont[j] == 1:
                MultiRes[j].append(f"f_{num} {file_name}")
            elif Multicont[j] > 1:
                MultiRes[j].append(f"f_{num} {file_name} (#: {Multicont[j]})")      
            

    con_beg = ['{ ' for i in range(len(Multiphrase))]
    con_end = [' }' for i in range(len(Multiphrase))]

    #display:
    print(f"Results: {sum([len(MultiRes_adi[k])for k in range(len(MultiRes_adi))])}")
    lis_disp = list(zip(con_beg, Multiphrase, con_end, [': ' for i in range(len(Multiphrase))], Multicont_total))

    salida = False
    while not salida:
        print('')
    ##    print('MENU')
        print('-'*54)
        print('(2)Ph context(all)  (3)Ph context(res)  (4)Exit')
        print('-'*54)
        print('option?:')
        opcion = int(input())

        if opcion == 1:
            for k in range(len(MultiRes)):
                if len(MultiRes[k]) > 0:
                    if len(Multiphrase[k])>40:
                        print('Content: { '+ Multiphrase[k][:40]+'.'*5 +' }' + addcase)
                    else:
                        print('Content: { '+ Multiphrase[k] +' }' + addcase)
                    print('Found in:')
                    print(*MultiRes[k], sep='\n')      

        elif opcion == 2:
            for k in range(len(MultiRes_adi)):
                if len(MultiRes_adi[k]) > 0:
                    print(*MultiRes_adi[k], sep='\n')
                    print('-'*10)

        elif opcion == 3:
            for k in range(len(MultiRes_adi)):
                if len(MultiRes_adi[k]) > 0:
                    MuRes_sel = [x[4:] for x in MultiRes_adi[k]]
                    MuRes_sel = sorted(list(set(MuRes_sel)))
                    print(*MuRes_sel, sep='\n')
                    print('-'*10)

        elif opcion == 4:
            print('done '+'/\\'*22)
            salida = True

#_________________________________________________
if Filenames_only and len(Res_namefs)>0:
    print('**** in file names only ****')
    print(*Res_namefs, sep='\n')

            

